#!/usr/bin/env python3
"""Render AI-authored ZStack fault report JSON into the standard DOCX template."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph


SKILL_DIR = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = SKILL_DIR / "assets" / "ZStack企业版-故障分析报告模板V1.2.docx"


INFO_FIELDS = [
    ("project_name", "项目名称"),
    ("customer_name", "客户名称"),
    ("fault_impact_scope", "故障影响和范围"),
    ("customer_contact", "客户联系人"),
    ("reporter", "故障报告人"),
    ("software_product", "软件产品"),
    ("version", "版本号"),
    ("fault_start_time", "故障发生时间"),
    ("business_recovery_time", "业务恢复时间"),
    ("fault_duration", "故障总耗时长"),
    ("fault_category", "故障类别"),
    ("responsible_department", "故障责任部门"),
    ("fault_level", "故障级别"),
]

TEXT_BLOCKS = {
    "analysis_process": "分析处理过程",
    "root_cause_analysis": "原因分析",
    "improvement_plan": "后续改进与预防方案",
}


def blank(value: Any) -> bool:
    return value is None or value == "" or value == []


def scalar(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        return "；".join(scalar(item) for item in value if scalar(item))
    if isinstance(value, dict):
        return "；".join(f"{key}：{scalar(item)}" for key, item in value.items() if not blank(item))
    return str(value)


def as_list(value: Any) -> list[Any]:
    if blank(value):
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, str):
        lines = [line.strip() for line in value.splitlines() if line.strip()]
        return lines or [value.strip()]
    return [value]


def format_items(value: Any) -> list[str]:
    lines: list[str] = []
    for idx, item in enumerate(as_list(value), 1):
        if isinstance(item, dict):
            parts = []
            for key, val in item.items():
                if not blank(val):
                    parts.append(f"{key}：{scalar(val)}")
            line = "；".join(parts) if parts else "待补充"
        else:
            line = scalar(item) or "待补充"
        if re.match(r"^\d+[.、]", line) or "：" in line[:20] or ":" in line[:20]:
            lines.append(line)
        else:
            lines.append(f"{idx}. {line}")
    return lines


def ensure_rfonts(run_or_style_element: Any, heading: bool = False) -> None:
    rpr = run_or_style_element.get_or_add_rPr() if hasattr(run_or_style_element, "get_or_add_rPr") else None
    if rpr is None:
        rpr = run_or_style_element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_or_style_element.append(rpr)
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    latin = "majorHAnsi" if heading else "minorHAnsi"
    east_asia = "majorEastAsia" if heading else "minorEastAsia"
    bidi = "majorBidi" if heading else "minorBidi"
    for attr, theme in [
        ("w:asciiTheme", latin),
        ("w:hAnsiTheme", latin),
        ("w:eastAsiaTheme", east_asia),
        ("w:cstheme", bidi),
    ]:
        rfonts.set(qn(attr), theme)


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def write_rich_text(paragraph: Paragraph, text: str, bold: bool = False) -> None:
    clear_paragraph(paragraph)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    heading = paragraph.style.name.startswith("Heading")
    segments = re.split(r"(\*\*.*?\*\*)", text)
    label_done = False
    for segment in segments:
        if not segment:
            continue
        segment_bold = bold
        if segment.startswith("**") and segment.endswith("**") and len(segment) >= 4:
            segment = segment[2:-2]
            segment_bold = True
        if not label_done and not segment_bold:
            match = re.match(r"^([^：:]{1,18}[：:])(.*)$", segment)
            if match:
                label_run = paragraph.add_run(match.group(1))
                label_run.bold = True
                ensure_rfonts(label_run._element, heading=heading)
                segment = match.group(2)
                label_done = True
        if segment:
            run = paragraph.add_run(segment)
            run.bold = segment_bold
            ensure_rfonts(run._element, heading=heading)


def set_paragraph_text(paragraph: Paragraph, text: str, bold: bool = False) -> None:
    write_rich_text(paragraph, text, bold=bold)


def set_cell_lines(cell: _Cell, lines: list[str], bold_first: bool = False) -> None:
    cell.text = ""
    if not lines:
        lines = [""]
    first = cell.paragraphs[0]
    write_rich_text(first, lines[0], bold=bold_first)
    for line in lines[1:]:
        paragraph = cell.add_paragraph()
        write_rich_text(paragraph, line)


def set_cell_text(cell: _Cell, text: str, bold: bool = False) -> None:
    set_cell_lines(cell, [text], bold_first=bold)


def table_by_first_cell(doc: Document, first_cell_text: str):
    for table in doc.tables:
        if table.rows and table.rows[0].cells and table.rows[0].cells[0].text.strip() == first_cell_text:
            return table
    raise ValueError(f"Template table not found: {first_cell_text}")


def update_cover(doc: Document, data: dict[str, Any]) -> None:
    replacements = {
        "云平台数据库残留XXXX": data.get("report_title"),
        "——问题故障报告": data.get("report_subtitle"),
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements and not blank(replacements[text]):
            set_paragraph_text(paragraph, scalar(replacements[text]), bold=text == "云平台数据库残留XXXX")


def find_label_cells(table, label: str) -> list[tuple[int, int]]:
    locations = []
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if cell.text.strip() == label:
                locations.append((row_index, col_index))
    return locations


def set_value_after_label(table, label: str, value: str) -> None:
    locations = find_label_cells(table, label)
    if not locations:
        return
    row_index, col_index = locations[0]
    row = table.rows[row_index]
    target_col = col_index + 1
    if target_col < len(row.cells):
        set_cell_text(row.cells[target_col], value)


def update_info_table(doc: Document, data: dict[str, Any]) -> None:
    table = table_by_first_cell(doc, "故障基本信息")
    for key, label in INFO_FIELDS:
        if key in data:
            set_value_after_label(table, label, scalar(data.get(key)))
    if "fault_description" in data:
        set_cell_lines(table.rows[11].cells[0], format_items(data.get("fault_description")))


def update_text_block(doc: Document, title: str, lines: list[str]) -> None:
    table = table_by_first_cell(doc, title)
    if len(table.rows) < 2:
        raise ValueError(f"Template text block has no content row: {title}")
    set_cell_lines(table.rows[1].cells[0], lines)


def update_text_blocks(doc: Document, data: dict[str, Any]) -> None:
    for key, title in TEXT_BLOCKS.items():
        if key in data:
            update_text_block(doc, title, format_items(data.get(key)))


def template_summary(doc: Document) -> dict[str, Any]:
    return {
        "cover_fields": ["report_title", "report_subtitle"],
        "info_fields": [{"key": key, "label": label} for key, label in INFO_FIELDS],
        "text_blocks": [{"key": key, "title": title} for key, title in TEXT_BLOCKS.items()],
    }


def iter_cell_paragraphs(cell: _Cell) -> Iterable[Paragraph]:
    for paragraph in cell.paragraphs:
        yield paragraph
    for table in cell.tables:
        for row in table.rows:
            for nested_cell in row.cells:
                yield from iter_cell_paragraphs(nested_cell)


def iter_all_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_cell_paragraphs(cell)


def apply_style_contract(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        ensure_rfonts(style._element, heading=style_name.startswith("Heading"))
        style.paragraph_format.line_spacing = 1.5
        if style_name == "Normal":
            style.paragraph_format.space_after = Pt(6)
            style.font.size = Pt(11)
        else:
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(6)
            style.font.bold = True

    for paragraph in iter_all_paragraphs(doc):
        paragraph.paragraph_format.line_spacing = 1.5
        for run in paragraph.runs:
            ensure_rfonts(run._element, heading=paragraph.style.name.startswith("Heading"))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.5
                    paragraph.paragraph_format.space_after = Pt(0)


def generate(data: dict[str, Any], template: Path, output: Path) -> None:
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(template))

    update_cover(doc, data)
    update_info_table(doc, data)
    update_text_blocks(doc, data)
    apply_style_contract(doc)

    if not blank(data.get("report_title")):
        doc.core_properties.title = scalar(data.get("report_title"))
    doc.core_properties.subject = "ZStack fault analysis report"
    doc.save(str(output))


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError("Input JSON must be an object.")
    return data


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_json", nargs="?", help="UTF-8 JSON input file.")
    parser.add_argument("--out", help="Output DOCX path.")
    parser.add_argument("--template", default=str(DEFAULT_TEMPLATE), help="Template DOCX path.")
    parser.add_argument("--print-template", action="store_true", help="Print template fields as JSON and exit.")
    args = parser.parse_args(argv)

    template = Path(args.template)
    if args.print_template:
        doc = Document(str(template))
        json.dump(template_summary(doc), sys.stdout, ensure_ascii=True, indent=2)
        sys.stdout.write("\n")
        return 0

    if not args.input_json or not args.out:
        parser.error("input_json and --out are required unless --print-template is used.")

    generate(load_json(Path(args.input_json)), template, Path(args.out))
    print(f"Generated: {Path(args.out).resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
