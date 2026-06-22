#!/usr/bin/env python3
"""Render AI-authored ZStack change proposal JSON into the standard DOCX template."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph


SKILL_DIR = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = SKILL_DIR / "assets" / "2021XX-XX项目XX问题-ZStack变更方案模板v1.2.docx"
CHECKLIST_TITLE = "ZStack云计算运维风险checklist"


def blank(value: Any) -> bool:
    return value is None or value == "" or value == []


def scalar(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        return "；".join(scalar(item) for item in value if scalar(item))
    if isinstance(value, dict):
        return "；".join(f"{key}：{scalar(item)}" for key, item in value.items() if not blank(item))
    return str(value)


def as_list(value: Any) -> list[Any]:
    if blank(value):
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, str):
        lines = [line.strip() for line in value.splitlines() if line.strip()]
        return lines or [value.strip()]
    return [value]


def format_items(value: Any) -> list[str]:
    lines: list[str] = []
    for idx, item in enumerate(as_list(value), 1):
        if isinstance(item, dict):
            parts = []
            for key, val in item.items():
                if not blank(val):
                    parts.append(f"{key}：{scalar(val)}")
            line = "；".join(parts) if parts else "待补充"
        else:
            line = scalar(item) or "待补充"
        if re.match(r"^\d+[.、]", line):
            lines.append(line)
        else:
            lines.append(f"{idx}. {line}")
    return lines


def normalize(value: str) -> str:
    return re.sub(r"\s+", "", value or "").lower()


def ensure_rfonts(run_or_style_element: Any, heading: bool = False) -> None:
    rpr = run_or_style_element.get_or_add_rPr() if hasattr(run_or_style_element, "get_or_add_rPr") else None
    if rpr is None:
        rpr = run_or_style_element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_or_style_element.append(rpr)
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    latin = "majorHAnsi" if heading else "minorHAnsi"
    east_asia = "majorEastAsia" if heading else "minorEastAsia"
    bidi = "majorBidi" if heading else "minorBidi"
    for attr, theme in [
        ("w:asciiTheme", latin),
        ("w:hAnsiTheme", latin),
        ("w:eastAsiaTheme", east_asia),
        ("w:cstheme", bidi),
    ]:
        rfonts.set(qn(attr), theme)


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def write_rich_text(paragraph: Paragraph, text: str, bold: bool = False) -> None:
    clear_paragraph(paragraph)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    heading = paragraph.style.name.startswith("Heading")
    segments = re.split(r"(\*\*.*?\*\*)", text)
    label_done = False
    for segment in segments:
        if not segment:
            continue
        segment_bold = bold
        if segment.startswith("**") and segment.endswith("**") and len(segment) >= 4:
            segment = segment[2:-2]
            segment_bold = True
        if not label_done and not segment_bold:
            match = re.match(r"^([^：:]{1,16}[：:])(.*)$", segment)
            if match:
                label_run = paragraph.add_run(match.group(1))
                label_run.bold = True
                ensure_rfonts(label_run._element, heading=heading)
                segment = match.group(2)
                label_done = True
        if segment:
            run = paragraph.add_run(segment)
            run.bold = segment_bold
            ensure_rfonts(run._element, heading=heading)


def set_paragraph_text(paragraph: Paragraph, text: str, bold: bool = False) -> None:
    write_rich_text(paragraph, text, bold=bold)


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    set_paragraph_text(new_para, text)
    return new_para


def insert_paragraph_before(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    set_paragraph_text(new_para, text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Template paragraph not found: {text}")


def remove_paragraphs_between(start: Paragraph, end: Paragraph | None, doc: Document) -> None:
    deleting = False
    targets: list[Paragraph] = []
    for paragraph in list(doc.paragraphs):
        if paragraph._element is start._element:
            deleting = True
            continue
        if end is not None and paragraph._element is end._element:
            break
        if deleting:
            targets.append(paragraph)
    for paragraph in targets:
        remove_paragraph(paragraph)


def replace_section(doc: Document, heading_text: str, next_heading_text: str | None, lines: list[str]) -> None:
    heading = find_paragraph(doc, heading_text)
    next_heading = find_paragraph(doc, next_heading_text) if next_heading_text else None
    remove_paragraphs_between(heading, next_heading, doc)
    anchor = heading
    for line in lines:
        anchor = insert_paragraph_after(anchor, line)


def replace_paragraphs_before(doc: Document, anchor_heading_text: str, previous_heading_text: str, lines: list[str]) -> None:
    previous = find_paragraph(doc, previous_heading_text)
    anchor = find_paragraph(doc, anchor_heading_text)
    remove_paragraphs_between(previous, anchor, doc)
    for line in lines:
        insert_paragraph_before(anchor, line)


def table_by_first_cell(doc: Document, first_cell_text: str):
    for table in doc.tables:
        if table.rows and table.rows[0].cells and table.rows[0].cells[0].text.strip() == first_cell_text:
            return table
    raise ValueError(f"Template table not found: {first_cell_text}")


def set_cell_text(cell: _Cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    write_rich_text(paragraph, text, bold=bold)


def update_cover(doc: Document, data: dict[str, Any]) -> None:
    replacements = {
        "ZStack运维变更方案": data.get("cover_title"),
        "XX变更": data.get("change_title") or data.get("change_subject"),
        "风险等级：X": f"风险等级：{scalar(data.get('risk_level'))}" if not blank(data.get("risk_level")) else None,
        "时间：XX.XX.XX": f"时间：{scalar(data.get('document_date'))}" if not blank(data.get("document_date")) else None,
    }
    for old, new in replacements.items():
        if not blank(new):
            set_paragraph_text(find_paragraph(doc, old), scalar(new), bold=old in {"ZStack运维变更方案", "风险等级：X"})


def update_config_table(doc: Document, data: dict[str, Any]) -> None:
    table = table_by_first_cell(doc, "软件信息")
    for row, key in zip(table.rows, ["software_info", "hardware_config", "business_info"]):
        if len(row.cells) >= 2 and key in data:
            set_cell_text(row.cells[1], scalar(data.get(key)))


def checklist_rows(doc: Document) -> list[dict[str, str]]:
    table = table_by_first_cell(doc, CHECKLIST_TITLE)
    rows: list[dict[str, str]] = []
    for index, row in enumerate(table.rows[2:], 2):
        cells = row.cells
        if len(cells) < 5:
            continue
        rows.append(
            {
                "row_index": str(index),
                "module": cells[0].text.strip(),
                "operation": cells[1].text.strip(),
                "impact": cells[2].text.strip(),
                "level": cells[3].text.strip(),
            }
        )
    return rows


def collect_checklist_items(data: dict[str, Any]) -> tuple[set[str], list[str]]:
    warnings: list[str] = []
    explicit: set[str] = set()

    for item in as_list(data.get("checklist_items")):
        value = item.get("operation") if isinstance(item, dict) else item
        if not blank(value):
            explicit.add(normalize(scalar(value)))

    for item in as_list(data.get("checklist_decisions")):
        if not isinstance(item, dict):
            continue
        involved = item.get("involved")
        if isinstance(involved, str):
            involved = involved.strip().lower() in {"true", "yes", "y", "1", "是", "涉及"}
        if involved:
            value = item.get("operation") or item.get("action")
            if not blank(value):
                explicit.add(normalize(scalar(value)))

    legacy_values = []
    legacy_values.extend(as_list(data.get("checklist_keywords")))
    for risk in as_list(data.get("risks")):
        if isinstance(risk, dict):
            legacy_values.extend(as_list(risk.get("checklist_keywords")))
    if legacy_values:
        warnings.append("checklist_keywords is ignored; use exact checklist_items or checklist_decisions.")

    return explicit, warnings


def update_risk_checklist(doc: Document, data: dict[str, Any]) -> list[str]:
    table = table_by_first_cell(doc, CHECKLIST_TITLE)
    explicit, warnings = collect_checklist_items(data)
    mark_unmatched = bool(data.get("mark_unmatched_checklist_as_no"))
    matched: set[str] = set()

    if not explicit and not mark_unmatched:
        return warnings

    for row in table.rows[2:]:
        cells = row.cells
        if len(cells) < 5:
            continue
        operation = normalize(cells[1].text.strip())
        involved = operation in explicit
        if involved:
            matched.add(operation)
            set_cell_text(cells[4], "是", bold=True)
        elif mark_unmatched:
            set_cell_text(cells[4], "否")

    unmatched = sorted(explicit - matched)
    for item in unmatched:
        warnings.append(f"checklist item not found in template operation column: {item}")
    return warnings


def update_risk_text(doc: Document, data: dict[str, Any]) -> None:
    if "risks" not in data:
        return
    lines = format_items(data.get("risks"))
    replace_paragraphs_before(doc, "3.2风险预案", "3.1风险清单", lines)


def update_change_plan(doc: Document, data: dict[str, Any]) -> None:
    if "change_plan" in data:
        replace_section(doc, "变更计划", None, format_items(data.get("change_plan")))
        return

    field_lines = []
    if "change_time" in data:
        field_lines.append(f"变更时间：{scalar(data.get('change_time'))}")
    if "maintenance_window" in data:
        field_lines.append(f"变更窗口：{scalar(data.get('maintenance_window'))}")
    if "executor" in data or "executor_phone" in data:
        field_lines.append(f"变更执行人：{scalar(data.get('executor'))} 联系电话 {scalar(data.get('executor_phone'))}")
    if "supervisor" in data or "supervisor_phone" in data:
        field_lines.append(f"变更监督人：{scalar(data.get('supervisor'))} 联系电话 {scalar(data.get('supervisor_phone'))}")
    if field_lines:
        replace_section(doc, "变更计划", None, field_lines)


def apply_section_updates(doc: Document, data: dict[str, Any]) -> None:
    section_map = [
        ("overview", "变更概述", "变更步骤"),
        ("change_principles", "2.2 变更原则及变更范围", "2.3 变更整体流程"),
        ("overall_flow", "2.3 变更整体流程", "2.4 变更具体步骤"),
        ("detailed_steps", "2.4 变更具体步骤", "风险评估"),
        ("risk_mitigations", "3.2风险预案", "回退方案"),
        ("rollback_plan", "回退方案", "紧急预案"),
        ("emergency_plan", "紧急预案", "变更计划"),
    ]
    for key, heading, next_heading in section_map:
        if key in data:
            replace_section(doc, heading, next_heading, format_items(data.get(key)))


def iter_cell_paragraphs(cell: _Cell) -> Iterable[Paragraph]:
    for paragraph in cell.paragraphs:
        yield paragraph
    for table in cell.tables:
        for row in table.rows:
            for nested_cell in row.cells:
                yield from iter_cell_paragraphs(nested_cell)


def iter_all_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_cell_paragraphs(cell)


def apply_style_contract(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        ensure_rfonts(style._element, heading=style_name.startswith("Heading"))
        style.paragraph_format.line_spacing = 1.5
        if style_name == "Normal":
            style.paragraph_format.space_after = Pt(6)
            style.font.size = Pt(11)
        else:
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(6)
            style.font.bold = True

    for paragraph in iter_all_paragraphs(doc):
        paragraph.paragraph_format.line_spacing = 1.5
        for run in paragraph.runs:
            ensure_rfonts(run._element, heading=paragraph.style.name.startswith("Heading"))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.5
                    paragraph.paragraph_format.space_after = Pt(0)


def generate(data: dict[str, Any], template: Path, output: Path) -> list[str]:
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(template))

    update_cover(doc, data)
    update_config_table(doc, data)
    warnings = update_risk_checklist(doc, data)
    apply_section_updates(doc, data)
    update_risk_text(doc, data)
    update_change_plan(doc, data)

    apply_style_contract(doc)
    if not blank(data.get("change_title")):
        doc.core_properties.title = scalar(data.get("change_title"))
    doc.core_properties.subject = "ZStack change proposal"
    doc.save(str(output))
    return warnings


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError("Input JSON must be an object.")
    return data


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_json", nargs="?", help="UTF-8 JSON input file.")
    parser.add_argument("--out", help="Output DOCX path.")
    parser.add_argument("--template", default=str(DEFAULT_TEMPLATE), help="Template DOCX path.")
    parser.add_argument("--print-checklist", action="store_true", help="Print template checklist rows as JSON and exit.")
    args = parser.parse_args(argv)

    template = Path(args.template)
    if args.print_checklist:
        doc = Document(str(template))
        json.dump(checklist_rows(doc), sys.stdout, ensure_ascii=True, indent=2)
        sys.stdout.write("\n")
        return 0

    if not args.input_json or not args.out:
        parser.error("input_json and --out are required unless --print-checklist is used.")

    warnings = generate(load_json(Path(args.input_json)), template, Path(args.out))
    for warning in warnings:
        print(f"WARNING: {warning}", file=sys.stderr)
    print(f"Generated: {Path(args.out).resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
